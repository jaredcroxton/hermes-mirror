#!/usr/bin/env python3
"""
Working template: generate a formatted Word document with sections.
Adjust the content blocks and output_path as needed.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# Default style
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.space_before = Pt(0)

# ---- Formatting helpers ----

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(18)
    run.bold = True

def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

def add_date(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

def add_section_header(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.bold = True

def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)

def add_final(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

# ---- Build document below this line ----

add_title("DOCUMENT TITLE")
add_subtitle("Subtitle or tagline goes here")
add_date("Organisation | Month YYYY")

add_section_header("FIRST SECTION HEADER")
add_body("Body paragraph text goes here.")
add_body("Second body paragraph.")

add_section_header("SECOND SECTION HEADER")
add_body("More body text.")

add_final("Closing line goes here.")

# ---- Save ----
output_path = "/Users/jc/Documents/Output Document.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
